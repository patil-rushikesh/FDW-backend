from flask import Flask, request, jsonify, send_file, make_response, send_from_directory
from flask_pymongo import PyMongo
from bson.json_util import dumps
import os
import bcrypt
from dotenv import load_dotenv
from mail import send_username_password_mail
from flask_cors import CORS  # Add this import
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import json
import requests
from docx import Document
from werkzeug.utils import secure_filename
from gridfs import GridFS
from bson.objectid import ObjectId
import math
# Add this import at the top\
from verification_commity import create_verification_blueprint
# Add this import at the top
from faculty_list import faculty_list
# Add this import at the top of app.py
from forgot_password import forgot_password



# Load environment variables
load_dotenv()

app = Flask(__name__)
# Configure CORS
CORS(app, resources={
    r"/*": {
        "origins": ["http://10.10.1.18:5173", "http://127.0.0.1:5173","http://localhost:5173"],  # Your React app's URLs
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "supports_credentials": True
    }
})

# MongoDB Configuration
app.config["MONGO_URI"] = os.getenv("MONGO_URI")
app.config["MONGO_URI_FDW"] = os.getenv("MONGO_URI_FDW")

mongo = PyMongo(app, uri=app.config["MONGO_URI"])
mongo_fdw = PyMongo(app, uri=app.config["MONGO_URI_FDW"])

# Collections
db_users = mongo.db.users
db_signin = mongo.db.signin

# GridFS instance
fs = GridFS(mongo_fdw.db)

# Department-based collections in the FDW database
department_collections = {
    "AIML": mongo_fdw.db.AIML,
    "ASH": mongo_fdw.db.ASH,
    "Civil": mongo_fdw.db.Civil,
    "Computer": mongo_fdw.db.Computer,
    "Computer(Regional)": mongo_fdw.db.Computer_Regional,
    "ENTC": mongo_fdw.db.ENTC,
    "IT": mongo_fdw.db.IT,
    "Mechanical": mongo_fdw.db.Mechanical
}

# Health check
@app.route('/', methods=['GET'])
def health_check():
    return jsonify({"message": "Welcome to FDW project"}), 200

# Create a new user
@app.route('/users', methods=['POST'])
def add_user():
    data = request.json
    required_fields = ["_id", "name", "role", "dept", "mail", "mob"]

    if not data or not all(k in data for k in required_fields):
        return jsonify({"error": "Missing required fields"}), 400
    
    if "desg" not in data:
        data["desg"] = "Faculty"
        
    if data['desg'] == 'Associate Dean':
        # Get the Dean's ID from the request data
        higher_dean_id = data.get('higherDean')
        if not higher_dean_id:
            return jsonify({"error": "Higher Dean ID is required for Associate Dean"}), 400
            
        try:
            # Check if lookup document exists
            lookup_doc = mongo.db.lookup.find_one({"_id": "deans"})
            
            if lookup_doc:
                # Update existing lookup document
                mongo.db.lookup.update_one(
                    {"_id": "deans"},
                    {"$push": {f"higherDeanId.{higher_dean_id}": data["_id"]}},
                    upsert=True
                )
            else:
                # Create new lookup document
                mongo.db.lookup.insert_one({
                    "_id": "deans",
                    "higherDeanId": {
                        higher_dean_id: [{"id" : data["_id"], "department" : data["dept"]}]
                    }
                })
        except Exception as e:
            return jsonify({"error": f"Error updating lookup collection: {str(e)}"}), 500

        

    # Add new default fields
    data["isInVerificationPanel"] = False
    data["facultyToVerify"] = {}

    try:
        # Insert into users collection
        db_users.insert_one(data)

        # Hash the password (use _id as the password initially)
        salt = bcrypt.gensalt()
        hashed_password = bcrypt.hashpw(data["_id"].encode('utf-8'), salt)

        # Insert into signin collection
        db_signin.insert_one({"_id": data["_id"], "password": hashed_password})

        # Get the correct department collection
        department = data["dept"]
        collection = department_collections.get(department)

        if collection is not None:
            # Update lookup document
            collection.update_one(
                {"_id": "lookup"},
                {"$set": {f"data.{data['_id']}": data["role"]}},
                upsert=True
            )

            # Create empty document for the user
            empty_doc = {
        "_id": data["_id"],
        "status": "pending",
        "isUpdated": False,
        "grand_total": {
            "grand_total": 0,
            "status": "pending"
        },
        "A": {
            "1": {
                "courses": {},
                "total_marks": 0
            },
            "2": {
                "courses": {},
                "semesterScores": {
                    "Sem I": 0,
                    "Sem II": 0
                },
                "total_marks": 0
            },
            "3": {
                "elearningInstances": 0,
                "total_marks": 0
            },
            "4": {
                "courses": {},
                "total_marks": 0
            },
            "5": {
                "weeklyLoadSem1": 0,
                "weeklyLoadSem2": 0,
                "adminResponsibility": 0,
                "cadre": data["role"],
                "total_marks": 0
            },
            "6": {
                "projectsGuided": 0,
                "total_marks": 0
            },
            "7": {
                "courses": {},
                "total_marks": 0
            },
            "8": {
                "ptgMeetings": 0,
                "total_marks": 0
            },
            "total_marks": 0
        },
        "B": {
            "1": {
                "journalPapers": {
                    "sciCount": 0, "sciProof": "", "ver_sciMarks": 0,
                    "esciCount": 0, "esciProof": "", "ver_esciMarks": 0,
                    "scopusCount": 0, "scopusProof": "", "ver_scopusMarks": 0,
                    "ugcCareCount": 0, "ugcCareProof": "", "ver_ugcCareMarks": 0,
                    "otherCount": 0, "otherProof": "", "ver_otherMarks": 0,
                    "marks": 0, "verified_marks": 0
                }
            },
            "2": {
                "conferencePapers": {
                    "scopusWosCount": 0, "scopusWosProof": "", "ver_scopusWosMarks": 0,
                    "otherCount": 0, "otherProof": "", "ver_otherMarks": 0,
                    "marks": 0, "verified_marks": 0
                }
            },
            '3': {
                'bookChapters': {
                    'scopusWosCount': 0, 'scopusWosProof': '', 'ver_scopusWosMarks': 0,
                    'otherCount': 0, 'otherProof': '', 'ver_otherMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '4': {
                'books': {
                    'scopusWosCount': 0, 'scopusWosProof': '', 'ver_scopusWosMarks': 0,
                    'nonIndexedCount': 0, 'nonIndexedProof': '', 'ver_nonIndexedMarks': 0,
                    'localCount': 0, 'localProof': '', 'ver_localMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '5': {
                'citations': {
                    'webOfScienceCount': 0, 'webOfScienceProof': '', 'ver_webOfScienceMarks': 0,
                    'scopusCount': 0, 'scopusProof': '', 'ver_scopusMarks': 0,
                    'googleScholarCount': 0, 'googleScholarProof': '', 'ver_googleScholarMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '6': {
                'copyrightIndividual': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '7': {
                'copyrightInstitute': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '8': {
                'patentIndividual': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'publishedCount': 0, 'publishedProof': '', 'ver_publishedMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '9': {
                'patentInstitute': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'publishedCount': 0, 'publishedProof': '', 'ver_publishedMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '10': {
                'researchGrants': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '11': {
                'trainingPrograms': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '12': {
                'nonResearchGrants': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '13': {
                'productDevelopment': {
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'developedCount': 0, 'developedProof': '', 'ver_developedMarks': 0,
                    'pocCount': 0, 'pocProof': '', 'ver_pocMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '14': {
                'startup': {
                    'revenueFiftyKCount': 0, 'revenueFiftyKProof': '', 'ver_revenueFiftyKMarks': 0,
                    'fundsFiveLakhsCount': 0, 'fundsFiveLakhsProof': '', 'ver_fundsFiveLakhsMarks': 0,
                    'productsCount': 0, 'productsProof': '', 'ver_productsMarks': 0,
                    'pocCount': 0, 'pocProof': '', 'ver_pocMarks': 0,
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '15': {
                'awardsAndFellowships': {
                    'internationalAwardsCount': 0, 'internationalAwardsProof': '', 'ver_internationalAwardsMarks': 0,
                    'governmentAwardsCount': 0, 'governmentAwardsProof': '', 'ver_governmentAwardsMarks': 0,
                    'nationalAwardsCount': 0, 'nationalAwardsProof': '', 'ver_nationalAwardsMarks': 0,
                    'internationalFellowshipsCount': 0, 'internationalFellowshipsProof': '', 'ver_internationalFellowshipsMarks': 0,
                    'nationalFellowshipsCount': 0, 'nationalFellowshipsProof': '', 'ver_nationalFellowshipsMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '16': {
                'industryInteraction': {
                    'moUsCount': 0, 'moUsProof': '', 'ver_moUsMarks': 0,
                    'collaborationCount': 0, 'collaborationProof': '', 'ver_collaborationMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '17': {
                'internshipPlacement': {
                    'offersCount': 0, 'offersProof': '', 'ver_offersMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            "total_marks": 0,
            "final_verified_marks": 0,
            "verifier_id": ""
        
        },
        "C": {
            "1": {
                "qualification": {
                    "pdfCompleted": False,
                    "pdfOngoing": False,
                    "phdAwarded": False,
                    "marks": 0
                }
            },
            "2": {
                "trainingAttended": {
                    "twoWeekProgram": 0,
                    "oneWeekProgram": 0,
                    "twoToFiveDayProgram": 0,
                    "oneDayProgram": 0,
                    "marks": 0
                }
            },
            "3": {
                "trainingOrganized": {
                    "twoWeekProgram": 0,
                    "oneWeekProgram": 0,
                    "twoToFiveDayProgram": 0,
                    "oneDayProgram": 0,
                    "marks": 0
                }
            },
            "4": {
                "phdGuided": {
                    "degreesAwarded": 0,
                    "thesisSubmitted": 0,
                    "scholarsGuiding": 0,
                    "marks": 0
                }
            },
            "total_marks": 0
        },
        "D": {
            "portfolioType": "",
            "selfAwardedMarks": 0,
            "deanMarks": 0,
            "hodMarks": 0,
            "isMarkHOD": False,
            "isMarkDean": False,
            "isAdministrativeRole": False,
            "administrativeRole": "",
            "adminSelfAwardedMarks": 0,
            "directorMarks": 0,
            "adminDeanMarks": 0,
            "instituteLevelPortfolio": "",
            "departmentLevelPortfolio": "",
            "total_marks": 0,
            "isFirstTime": True
        },
        "E": {
            "total_marks": 0,
            "bullet_points": [],
            "verified_marks": 0,
            "verifier_comments": "",
            "isVerified": False
        }
    }
            collection.insert_one(empty_doc)
            mail_sent = send_username_password_mail(data["mail"], data["_id"], data["_id"],data["name"])
            if mail_sent:
                return jsonify({"message": f"User added successfully to {department}"}), 201
            else:
                return jsonify({"error": "Failed to send email notification"}), 500
        else:
            return jsonify({"error": "Invalid department"}), 400

    except Exception as e:
        print(str(e))
        return jsonify({"error": str(e)}), 500

# Get all users
@app.route('/users', methods=['GET'])
def get_users():
    users = db_users.find()
    return dumps(users), 200

# Get a user by ID
@app.route('/users/<string:user_id>', methods=['GET'])
def get_user(user_id):
    user = db_users.find_one({"_id": user_id})
    if user:
        return dumps(user), 200
    return jsonify({"error": "User not found"}), 404

# Update a user by ID
@app.route('/users/<string:user_id>', methods=['PUT'])
def update_user(user_id):
    data = request.json
    allowed_fields = ["name", "role", "dept", "mail", "mob","desg"]
    updated_data = {k: v for k, v in data.items() if k in allowed_fields}

    if not updated_data:
        return jsonify({"error": "No valid fields to update"}), 400

    result = db_users.update_one({"_id": user_id}, {"$set": updated_data})

    if result.modified_count:
        return jsonify({"message": "User updated successfully"}), 200
    return jsonify({"error": "User not found or no changes made"}), 404

# Delete a user by ID
@app.route('/users/<string:user_id>', methods=['DELETE'])
def delete_user(user_id):
    result = db_users.delete_one({"_id": user_id})
    db_signin.delete_one({"_id": user_id})  # Remove from signin collection
    
    if result.deleted_count:
        return jsonify({"message": "User deleted successfully"}), 200
    return jsonify({"error": "User not found"}), 404

# Migrate all users to signin collection
@app.route('/migrate_users', methods=['POST'])
def migrate_users():
    users = db_users.find()
    migrated_count = 0

    for user in users:
        user_id = user["_id"]
        if not db_signin.find_one({"_id": user_id}):
            salt = bcrypt.gensalt()
            hashed_password = bcrypt.hashpw(user_id.encode('utf-8'), salt)
            db_signin.insert_one({"_id": user_id, "password": hashed_password})
            migrated_count += 1
        department = user["dept"]
        collection = department_collections.get(department)
        if collection is not None:
            collection.update_one(
                {"_id": "lookup"},
                {"$set": {f"data.{user_id}": user["role"]}},
                upsert=True
            )
            

    return jsonify({"message": f"Migrated {migrated_count} users to signin collection"}), 200

# User login
@app.route('/login', methods=['POST', 'OPTIONS'])
def login():
    # Handle preflight request
    if request.method == 'OPTIONS':
        response = app.make_default_options_response()
        return response

    data = request.json
    if not data or not all(k in data for k in ["_id", "password"]):
        return jsonify({"error": "Missing required fields"}), 400

    user = db_signin.find_one({"_id": data["_id"]})
    if user and bcrypt.checkpw(data["password"].encode('utf-8'), user["password"]):
        user_data = db_users.find_one({"_id": data["_id"]})
        if not user_data:
            return jsonify({"error": "User data not found"}), 404

        # Create response data based on user type
        response_data = {
            "_id": user_data["_id"],
            "name": user_data.get("full_name") if user_data.get("isExternal") else user_data.get("name"),
            "role": user_data.get("role"),
            "dept": user_data.get("dept"),
            "isExternal": user_data.get("isExternal", False),
            "mail": user_data.get("mail"),
            "desg": user_data.get("desg", "Faculty"),
            "isAddedForInteraction" : user_data.get("isAddedForInteraction", False),
            "interactionDepartments" : user_data.get("interactionDepartments", []),
        }

        # Add external-specific fields if user is external
        if user_data.get("isExternal"):
            response_data.update({
                "specialization": user_data.get("specialization"),
                "organization": user_data.get("organization"),
                "facultyToReview": user_data.get("facultyToReview", []),
                "mob": user_data.get("mob"),
            })
        else:
            # Add regular faculty fields
            response_data.update({
                "mail": user_data.get("mail"),
                "mob": user_data.get("mob"),
                "isInVerificationPanel": user_data.get("isInVerificationPanel", False),
                "facultyToVerify": user_data.get("facultyToVerify", {})
            })

        response = app.response_class(
            response=dumps(response_data),
            status=200,
            mimetype='application/json'
        )
        return response

    return jsonify({"error": "Invalid credentials"}), 401



#Section Data Adding Start here
def calculate_grand_total(data):
    """Calculate grand total from all sections and determine status"""
    try:
        grand_total = 0
        form_status = "pending"  # Default status

        # Add Section A total if exists
        if 'A' in data and 'total_marks' in data['A']:
            grand_total += float(data['A']['total_marks'])

        # Add Section B total if exists
        if 'B' in data and 'total_marks' in data['B']:
            grand_total += float(data['B']['total_marks'])

        # Add Section C total if exists
        if 'C' in data and 'total_marks' in data['C']:
            grand_total += float(data['C']['total_marks'])

        # Add Section D total if exists
        if 'D' in data and 'total_marks' in data['D']:
            grand_total += float(data['D']['total_marks'])

        # Add Section E total if exists
        if 'E' in data and 'total_marks' in data['E']:
            grand_total += float(data['E']['total_marks'])

        # Return as dict with proper structure
        return {
            'grand_total': grand_total,
            'status': form_status
        }

    except Exception as e:
        print(f"Error calculating grand total: {str(e)}")
        return {
            'grand_total': 0,
            'status': "error"
        }

# Modify section handlers to update grand total
@app.route('/<department>/<user_id>/A', methods=['POST'])
def handle_post_A(department, user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400

        collection = department_collections.get(department)
        
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400
        
        lookup = collection.find_one({"_id": "lookup"}).get("data")
        if lookup is None:
            return jsonify({"error": "Invalid department"}), 400
        user = lookup.get(user_id)
        if user is None:
            return jsonify({"error": "Invalid user"}), 400
        print(data)
        # Update the document for the given user_id
        if 'total_marks' not in data['1']:
            data['1']['total_marks'] = 0
        if 'total_marks' not in data['2']:
            data['2']['total_marks'] = 0
        if 'total_marks' not in data['3']:
            data['3']['total_marks'] = 0
        if 'total_marks' not in data['4']:
            data['4']['total_marks'] = 0
        if 'total_marks' not in data['5']:
            data['5']['total_marks'] = 0
        if 'total_marks' not in data['6']:
            data['6']['total_marks'] = 0
        if 'total_marks' not in data['7']:
            data['7']['total_marks'] = 0
        if 'total_marks' not in data['8']:
            data['8']['total_marks'] = 0
        if 'total_marks' not in data:
            data['total_marks'] = 0
        
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {
                "A": data,
                "isUpdated": True,
                "status": "pending"  # Set initial status
            }},
            upsert=True
        )

        # Get updated document and calculate grand total
        updated_doc = collection.find_one({"_id": user_id})
        calculated_data = calculate_grand_total(updated_doc)

        # Update grand total and status
        collection.update_one(
            {"_id": user_id},
            {"$set": {
                "grand_total": calculated_data['grand_total'],
                "status": calculated_data['status']
            }}
        )

        if result.matched_count > 0:
            message = "Data updated successfully"
        else:
            message = "Data inserted successfully"

        return jsonify({
            "message": message,
            "grand_total": calculated_data['grand_total'],
            "status": calculated_data['status']
        }), 200

    except Exception as e:
        print(f"Error updating section A: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/A', methods=['GET'])
def get_section_A(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is not None:
            user = collection.find_one({"_id": user_id})
            if user:
                return jsonify(user.get("A"))
            return jsonify({"error": "User not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500
        
        

#Section B Data Adding Start here
@app.route('/<department>/<user_id>/B', methods=['POST'])
def handle_post_B(department, user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400
        
        # Access the collection named after the department
        collection = department_collections.get(department)
        
        #first we have to verify from lookup that the user exist in that department or not then only add data
        
        
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400
        
        lookup = collection.find_one({"_id": "lookup"}).get("data")
        print(lookup)
        if lookup is None:
            return jsonify({"error": "Invalid department"}), 400
        user = lookup.get(user_id)
        if user is None:
            return jsonify({"error": "Invalid user"}), 400
        
        
        
         # Initialize default structure for section B if not present
        

        # Initialize subsections with default values
        sections = {
            '1': {
                'journalPapers': {
                    'sciCount': 0, 'sciProof': '', 'ver_sciMarks': 0,
                    'esciCount': 0, 'esciProof': '', 'ver_esciMarks': 0,
                    'scopusCount': 0, 'scopusProof': '', 'ver_scopusMarks': 0,
                    'ugcCareCount': 0, 'ugcCareProof': '', 'ver_ugcCareMarks': 0,
                    'otherCount': 0, 'otherProof': '', 'ver_otherMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '2': {
                'conferencePapers': {
                    'scopusWosCount': 0, 'scopusWosProof': '', 'ver_scopusWosMarks': 0,
                    'otherCount': 0, 'otherProof': '', 'ver_otherMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '3': {
                'bookChapters': {
                    'scopusWosCount': 0, 'scopusWosProof': '', 'ver_scopusWosMarks': 0,
                    'otherCount': 0, 'otherProof': '', 'ver_otherMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '4': {
                'books': {
                    'scopusWosCount': 0, 'scopusWosProof': '', 'ver_scopusWosMarks': 0,
                    'nonIndexedCount': 0, 'nonIndexedProof': '', 'ver_nonIndexedMarks': 0,
                    'localCount': 0, 'localProof': '', 'ver_localMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '5': {
                'citations': {
                    'webOfScienceCount': 0, 'webOfScienceProof': '', 'ver_webOfScienceMarks': 0,
                    'scopusCount': 0, 'scopusProof': '', 'ver_scopusMarks': 0,
                    'googleScholarCount': 0, 'googleScholarProof': '', 'ver_googleScholarMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '6': {
                'copyrightIndividual': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '7': {
                'copyrightInstitute': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '8': {
                'patentIndividual': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'publishedCount': 0, 'publishedProof': '', 'ver_publishedMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '9': {
                'patentInstitute': {
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'publishedCount': 0, 'publishedProof': '', 'ver_publishedMarks': 0,
                    'grantedCount': 0, 'grantedProof': '', 'ver_grantedMarks': 0,
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '10': {
                'researchGrants': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '11': {
                'trainingPrograms': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '12': {
                'nonResearchGrants': {
                    'amount': 0, 'proof': '', 'ver_amountMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '13': {
                'productDevelopment': {
                    'commercializedCount': 0, 'commercializedProof': '', 'ver_commercializedMarks': 0,
                    'developedCount': 0, 'developedProof': '', 'ver_developedMarks': 0,
                    'pocCount': 0, 'pocProof': '', 'ver_pocMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '14': {
                'startup': {
                    'revenueFiftyKCount': 0, 'revenueFiftyKProof': '', 'ver_revenueFiftyKMarks': 0,
                    'fundsFiveLakhsCount': 0, 'fundsFiveLakhsProof': '', 'ver_fundsFiveLakhsMarks': 0,
                    'productsCount': 0, 'productsProof': '', 'ver_productsMarks': 0,
                    'pocCount': 0, 'pocProof': '', 'ver_pocMarks': 0,
                    'registeredCount': 0, 'registeredProof': '', 'ver_registeredMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '15': {
                'awardsAndFellowships': {
                    'internationalAwardsCount': 0, 'internationalAwardsProof': '', 'ver_internationalAwardsMarks': 0,
                    'governmentAwardsCount': 0, 'governmentAwardsProof': '', 'ver_governmentAwardsMarks': 0,
                    'nationalAwardsCount': 0, 'nationalAwardsProof': '', 'ver_nationalAwardsMarks': 0,
                    'internationalFellowshipsCount': 0, 'internationalFellowshipsProof': '', 'ver_internationalFellowshipsMarks': 0,
                    'nationalFellowshipsCount': 0, 'nationalFellowshipsProof': '', 'ver_nationalFellowshipsMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '16': {
                'industryInteraction': {
                    'moUsCount': 0, 'moUsProof': '', 'ver_moUsMarks': 0,
                    'collaborationCount': 0, 'collaborationProof': '', 'ver_collaborationMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            '17': {
                'internshipPlacement': {
                    'offersCount': 0, 'offersProof': '', 'ver_offersMarks': 0,
                    'marks': 0, 'verified_marks': 0
                }
            },
            
        }

        # Merge incoming data with default values
        for section, default_data in sections.items():
            if section not in data:
                print(f"Adding default value for {section}")
                data[section] = default_data
            else:
                for category, category_data in default_data.items():
                    if category not in data[section]:
                        print(f"Adding default value for {section} -> {category}")
                        data[section][category] = category_data
                    else:
                        for field, default_value in category_data.items():
                            if field not in data[section][category]:
                                print(f"Adding default value for {section} - {category} - {field}")
                                data[section][category][field] = default_value
        checkData = {"total_marks": 0,
            "final_verified_marks": 0,
            "verifier_id": ""}
        for field,value in checkData.items() :
            if field not in data:
                print(f"{field} is not present")
                data[field] = value
        # Rest of your existing code for database update
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        result = collection.update_one(
            {"_id": user_id},
            {"$set": {
                "B": data,
                "isUpdated": True
            }},
            upsert=True
        )
        print('added data in B')

        # Get updated document and calculate grand total
        updated_doc = collection.find_one({"_id": user_id})
        grand_total = calculate_grand_total(updated_doc)

        # Update grand total
        collection.update_one(
            {"_id": user_id},
            {"$set": {"grand_total": grand_total}}
        )
        print('grand total updated')
        
        if result.matched_count > 0:
            message = "Data updated successfully"
        else:
            message = "Data inserted successfully"
        
        return jsonify({
            "message": message,
            "grand_total": grand_total
        }), 200
    
    except Exception as e:
        print(f"Error updating section B: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/B', methods=['GET'])
def get_section_B(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is not None:
            user = collection.find_one({"_id": user_id})
            if user:
                return jsonify(user.get("B"))
            return jsonify({"error": "User not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

#Section C Data Adding Start here
@app.route('/<department>/<user_id>/C', methods = ['POST'])
def handle_post_C(department, user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400

        # Initialize default structure for section C if not present
        

        # Initialize subsections with default values
        sections = {
            '1': {
                'qualification': {
                    'pdfCompleted': False,
                    'pdfOngoing': False,
                    'phdAwarded': False,
                    'marks': 0
                }
            },
            '2': {
                'trainingAttended': {
                    'twoWeekProgram': 0,
                    'oneWeekProgram': 0,
                    'twoToFiveDayProgram': 0,
                    'oneDayProgram': 0,
                    'marks': 0
                }
            },
            '3': {
                'trainingOrganized': {
                    'twoWeekProgram': 0,
                    'oneWeekProgram': 0,
                    'twoToFiveDayProgram': 0,
                    'oneDayProgram': 0,
                    'marks': 0
                }
            },
            '4': {
                'phdGuided': {
                    'degreesAwarded': 0,
                    'thesisSubmitted': 0,
                    'scholarsGuiding': 0,
                    'marks': 0
                }
            },
           
        }

        # Merge incoming data with default values
        for section, default_data in sections.items():
            if section not in data:
                data[section] = default_data
            else:
                for category, category_data in default_data.items():
                    if isinstance(category_data, dict):
                        if category not in data[section]:
                            data[section][category] = category_data
                        else:
                            for field, default_value in category_data.items():
                                if field not in data[section][category]:
                                    data[section][category][field] = default_value
                    else:
                        if category not in data[section]:
                            data[section][category] = category_data
                            
        if  'total_marks' not in data:
            data['total_marks'] = 0

        # Update document with merged data
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        result = collection.update_one(
            {"_id": user_id},
            {"$set": {
                "C": data,
                "isUpdated": True
            }},
            upsert=True
        )

        # Update grand total
        updated_doc = collection.find_one({"_id": user_id})
        grand_total = calculate_grand_total(updated_doc)
        collection.update_one(
            {"_id": user_id},
            {"$set": {"grand_total": grand_total}}
        )

        return jsonify({
            "message": "Data updated successfully" if result.matched_count > 0 else "Data inserted successfully",
            "grand_total": grand_total
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/C', methods=['GET'])
def get_section_C(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is not None:
            user = collection.find_one({"_id": user_id})
            if user:
                return jsonify(user.get("C"))
            return jsonify({"error": "User not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500
                
#Section B Data Adding Start here
@app.route('/<department>/<user_id>/D', methods=['POST'])
def handle_post_D(department, user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400

        # Initialize default structure for section D if not present
        if 'D' not in data:
            data['D'] = {}

        # Initialize with default values
        default_data = {
            'portfolioType': '',
            'selfAwardedMarks': 0,
            'deanMarks': 0,
            'hodMarks': 0,
            'isMarkHOD': False,
            'isMarkDean': False,
            'isAdministrativeRole': False,
            'administrativeRole': '',
            'adminSelfAwardedMarks': 0,
            'directorMarks': 0,
            'adminDeanMarks': 0,
            'instituteLevelPortfolio': '',
            'departmentLevelPortfolio': '',
            'total_marks': 0,
            'isFirstTime': True
        }

        # Merge incoming data with default values
        for field, default_value in default_data.items():
            if field not in data['D']:
                data['D'][field] = default_value

        # Update document with merged data
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        result = collection.update_one(
            {"_id": user_id},
            {"$set": {
                "D": data['D'],
                "isUpdated": True
            }},
            upsert=True
        )

        # Update grand total
        updated_doc = collection.find_one({"_id": user_id})
        grand_total = calculate_grand_total(updated_doc)
        collection.update_one(
            {"_id": user_id},
            {"$set": {"grand_total": grand_total}}
        )

        return jsonify({
            "message": "Data updated successfully" if result.matched_count > 0 else "Data inserted successfully",
            "grand_total": grand_total
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/D', methods=['GET'])
def get_section_D(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is not None:
            user = collection.find_one({"_id": user_id})
            if user:
                return jsonify(user.get("D"))
            return jsonify({"error": "User not found"}), 404
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/total', methods=['GET'])
def get_grand_total(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        grand_total = calculate_grand_total(user_doc)
        
        return jsonify({
            "user_id": user_id,
            "grand_total": grand_total
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@app.route('/<department>/<user_id>/status', methods=['GET'])
def get_form_status(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        status = user_doc.get('status', 'pending')
        return jsonify({
            "user_id": user_id,
            "status": status
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
def fill_template_document(data, user_id, department):
    try:
        # Get user details directly from MongoDB
        user_data = db_users.find_one({"_id": user_id})
        if not user_data:
            raise Exception("User not found")

        # Load the template document
        doc = Document("Faculty Self Appraisal Scheme -PCCoE-24-25.docx")
        
        # Initial placeholders with user details
        placeholders = {
            '{faculty_name}': user_data.get('name', ''),
            '{faculty_designation}': user_data.get('role', ''),
            '{faculty_department}': user_data.get('dept', ''),
        }
        role = user_data.get('role', '')
        desg = user_data.get('desg', '')
        # Add all other placeholders using the data parameter
        
            # Rest of your existing code...
        section_a_marks = data.get('A', {}).get('total_marks', 0)
        Prof_A = 0
        Assoc_A = 0
        Assis_A = 0
        Prof_A_total_marks = 0
        Assoc_A_total_marks = 0
        Assis_A_total_marks = 0
        if role == 'Assistant Professor':
            section_a_marks = section_a_marks / 1.0
            Assis_A =  section_a_marks
            Assis_A_total_marks = (data.get('A', {}).get('total_marks', 0))
            
        if role == 'Associate Professor':
            section_a_marks = section_a_marks / 0.818
            Assoc_A =  section_a_marks
            Assoc_A_total_marks = (data.get('A', {}).get('total_marks', 0))
        elif role == 'Professor':
            section_a_marks = section_a_marks / 0.68        
            Prof_A =  section_a_marks
            Prof_A_total_marks = (data.get('A', {}).get('total_marks', 0))

        #adding all the marks in the B section and store in the variable
        b_total_verified = data['B']['1']['journalPapers']['verified_marks'] + data['B']['2']['conferencePapers']['verified_marks'] + data['B']['3']['bookChapters']['verified_marks'] + data['B']['4']['books']['verified_marks'] + data['B']['5']['citations']['verified_marks'] + data['B']['6']['copyrightIndividual']['verified_marks'] + data['B']['7']['copyrightInstitute']['verified_marks'] + data['B']['8']['patentIndividual']['verified_marks'] + data['B']['9']['patentInstitute']['verified_marks'] + data['B']['10']['researchGrants']['verified_marks'] + data['B']['11']['trainingPrograms']['verified_marks'] + data['B']['12']['nonResearchGrants']['verified_marks'] + data['B']['13']['productDevelopment']['verified_marks'] + data['B']['14']['startup']['verified_marks'] + data['B']['15']['awardsAndFellowships']['verified_marks'] + data['B']['16']['industryInteraction']['verified_marks'] + data['B']['17']['internshipPlacement']['verified_marks']
        # b_total = data['B']['1']['journalPapers']['total_marks'] + data['B']['2']['conferencePapers']['total_marks'] + data['B']['3']['bookChapters']['total_marks'] + data['B']['4']['books']['total_marks'] + data['B']['5']['citations']['total_marks'] + data['B']['6']['copyrightIndividual']['total_marks'] + data['B']['7']['copyrightInstitute']['total_marks'] + data['B']['8']['patentIndividual']['total_marks'] + data['B']['9']['patentInstitute']['total_marks'] + data['B']['10']['researchGrants']['total_marks'] + data['B']['11']['trainingPrograms']['total_marks'] + data['B']['12']['nonResearchGrants']['total_marks'] + data['B']['13']['productDevelopment']['total_marks'] + data['B']['14']['startup']['total_marks']
        b_total = data['B']['1']['journalPapers']['marks'] + data['B']['2']['conferencePapers']['marks'] + data['B']['3']['bookChapters']['marks'] + data['B']['4']['books']['marks'] + data['B']['5']['citations']['marks'] + data['B']['6']['copyrightIndividual']['marks'] + data['B']['7']['copyrightInstitute']['marks'] + data['B']['8']['patentIndividual']['marks'] + data['B']['9']['patentInstitute']['marks'] + data['B']['10']['researchGrants']['marks'] + data['B']['11']['trainingPrograms']['marks'] + data['B']['12']['nonResearchGrants']['marks'] + data['B']['13']['productDevelopment']['marks'] + data['B']['14']['startup']['marks'] + data['B']['15']['awardsAndFellowships']['marks'] + data['B']['16']['industryInteraction']['marks'] + data['B']['17']['internshipPlacement']['marks']
        print(f'#############{str(b_total)}##############{str(b_total_verified)}')
        Prof_B = 0
        Assoc_B = 0
        Assis_B = 0
        
        Prof_B_total_marks = 0
        Assoc_B_total_marks = 0
        Assis_B_total_marks = 0
        
        Prof_B_total_verified = 0
        Assoc_B_total_verified = 0
        Assis_B_total_verified = 0
        
        if role == 'Assistant Professor':
            Assis_B =  b_total
            Assis_B_total_marks = data['B']['total_marks']
            Assis_B_total_verified = data['B']['final_verified_marks']
        if role == 'Associate Professor':
            Assoc_B =  b_total
            Assoc_B_total_marks = data['B']['total_marks']
            Assoc_B_total_verified = data['B']['final_verified_marks']
        elif role == 'Professor':
            Prof_B =  b_total
            Prof_B_total_marks = data['B']['total_marks']
            Prof_B_total_verified = data['B']['final_verified_marks']
            
        #get the verifier name from the verifier id
        verifier_id = data['B']['verifier_id']
        verifier_name = 'Not Verified Yet'
        if verifier_id != '':
            verifier_name = db_users.find_one({"_id": verifier_id})['name']
        
        Prof_qualification_marks = 0
        qualification_marks  = 0
        if role == 'Assistant Professor':
            qualification_marks =  data['C']['1']['qualification']['marks']
        else:
            Prof_qualification_marks =  data['C']['1']['qualification']['marks']
            
        c_total = data['C']['1']['qualification']['marks'] + data['C']['2']['trainingAttended']['marks'] + data['C']['3']['trainingOrganized']['marks'] + data['C']['4']['phdGuided']['marks']
        Prof_C = 0
        Assoc_C = 0
        Assis_C = 0
        Prof_C_total_marks = 0
        Assoc_C_total_marks = 0
        Assis_C_total_marks = 0
        if role == 'Assistant Professor':
            Assis_C =  c_total
            Assis_C_total_marks = data['C']['total_marks']
        if role == 'Associate Professor':
            Assoc_C =  c_total
            Assoc_C_total_marks = data['C']['total_marks']
        elif role == 'Professor':
            Prof_C =  c_total
            Prof_C_total_marks = data['C']['total_marks']
        # Placeholders and their corresponding values from different sections
        self_awarded_marks = data['D']['selfAwardedMarks']
        hod_marks = data['D']['hodMarks']
        total_marks_D= data['D']['total_marks']
        assTotalMarks = 0
        assDeanHODMarks = 0
        assDeanDeanMarks = 0
        sumMarks_hod_dean = 0
        assSelfawardedmarks = 0
        if desg == 'Associate Dean' : 
            self_awarded_marks = 0
            total_marks_D = 0
            hod_marks = 0
            assDeanHODMarks = data['D']['hodMarks']
            assDeanDeanMarks = data['D']['deanMarks']
            sumMarks_hod_dean = (data['D']['hodMarks'] + data['D']['deanMarks']) / 2
            assSelfawardedmarks = data['D']['selfAwardedMarks']
            assTotalMarks = assSelfawardedmarks + sumMarks_hod_dean
        
        
        
        print('-----------before placeholders update-----------')
        
        placeholders.update({
            # Section A placeholders
            '{result_analysis_marks}': str(round(data['A']['1']['total_marks'], 2)),
            '{course_outcome_marks}': str(round(data['A']['2']['total_marks'], 2)),
            '{elearning_content_marks}': str(round(data['A']['3']['total_marks'], 2)),
            '{academic_engagement_marks}': str(round(data['A']['4']['total_marks'], 2)),
            '{teaching_load_marks}': str(round(data['A']['5']['total_marks'], 2)),
            '{projects_guided_marks}': str(round(data['A']['6']['total_marks'], 2)),
            '{student_feedback_marks}': str(round(data['A']['7']['total_marks'], 2)),
            '{ptg_meetings_marks}': str(round(data['A']['8']['total_marks'], 2)),
            '{section_a_total}': str(round(section_a_marks)),
            '{Prof_A}': str(round(Prof_A)),
            '{Assoc_A}': str(round(Assoc_A)),
            '{Assis_A}': str(round(Assis_A)),
            '{Prof_A_total_marks}': str(round(Prof_A_total_marks)),
            '{Assoc_A_total_marks}': str(round(Assoc_A_total_marks)),
            '{Assis_A_total_marks}': str(round(Assis_A_total_marks)),
            
            # Section B detailed placeholders - Updated to include verification marks
            # 1. Journal Papers
            
            
            
            '{sci_papers_marks}': str(data['B']['1']['journalPapers']['sciCount'] * 100),
            '{sci_papers_verified_marks}': str(data['B']['1']['journalPapers']['ver_sciMarks']),
            
            '{esci_papers_marks}': str(data['B']['1']['journalPapers']['esciCount'] * 50),
            '{esci_papers_verified_marks}': str(data['B']['1']['journalPapers']['ver_esciMarks']),
            
            '{scopus_papers_marks}': str(data['B']['1']['journalPapers']['scopusCount'] * 50),
            '{scopus_papers_verified_marks}': str(data['B']['1']['journalPapers']['ver_scopusMarks']),
            
            '{ugc_papers_marks}': str(data['B']['1']['journalPapers']['ugcCareCount'] * 10),
            '{ugc_papers_verified_marks}': str(data['B']['1']['journalPapers']['ver_ugcCareMarks']),
            
            '{other_papers_marks}': str(data['B']['1']['journalPapers']['otherCount'] * 5),
            '{other_papers_verified_marks}': str(data['B']['1']['journalPapers']['ver_otherMarks']),
            '{papers_published_marks}': str(data['B']['1']['journalPapers']['verified_marks']),
            
            # 2. Conferences
            
            '{scopus_conf_marks}': str(data['B']['2']['conferencePapers']['scopusWosCount'] * 30),
            '{scopus_conf_verified_marks}': str(data['B']['2']['conferencePapers']['ver_scopusWosMarks']),
            
            '{other_conf_marks}': str(data['B']['2']['conferencePapers']['otherCount'] * 5),
            '{other_conf_verified_marks}': str(data['B']['2']['conferencePapers']['ver_otherMarks']),
            '{conferences_marks}': str(data['B']['2']['conferencePapers']['verified_marks']),
            
            # 3. Book Chapters
            
            '{scopus_chapter_marks}': str(data['B']['3']['bookChapters']['scopusWosCount'] * 30),
            '{scopus_chapter_verified_marks}': str(data['B']['3']['bookChapters']['ver_scopusWosMarks']),
            
            '{other_chapter_marks}': str(data['B']['3']['bookChapters']['otherCount'] * 5),
            '{other_chapter_verified_marks}': str(data['B']['3']['bookChapters']['ver_otherMarks']),
            '{book_chapters_marks}': str(data['B']['3']['bookChapters']['verified_marks']),
            
            # 4. Books
            
            '{scopus_books_marks}': str(data['B']['4']['books']['scopusWosCount'] * 100),
            '{scopus_books_verified_marks}': str(data['B']['4']['books']['ver_scopusWosMarks']),
            
            '{national_books_marks}': str(data['B']['4']['books']['nonIndexedCount'] * 30),
            '{national_books_verified_marks}': str(data['B']['4']['books']['ver_nonIndexedMarks']),
            
            '{local_books_marks}': str(data['B']['4']['books']['localCount'] * 10),
            '{local_books_verified_marks}': str(data['B']['4']['books']['ver_localMarks']),
            '{books_marks}': str(data['B']['4']['books']['verified_marks']),
            
            # 5. Citations
            '{wos_citations_marks}': str(math.floor(data['B']['5']['citations']['webOfScienceCount'] / 3) * 3),
            '{wos_citations_verified_marks}': str(data['B']['5']['citations']['ver_webOfScienceMarks']),
            
            '{scopus_citations_marks}': str(math.floor(data['B']['5']['citations']['scopusCount'] / 3) * 3),
            '{scopus_citations_verified_marks}': str(data['B']['5']['citations']['ver_scopusMarks']),
            
            '{google_citations_marks}': str(math.floor(data['B']['5']['citations']['googleScholarCount'] / 3) * 1),
            '{google_citations_verified_marks}': str(data['B']['5']['citations']['ver_googleScholarMarks']),
            '{citations_marks}': str(data['B']['5']['citations']['verified_marks']),
            
            # 6. Copyright Individual
            
            '{individual_copyright_registered_marks}': str(data['B']['6']['copyrightIndividual']['registeredCount'] * 20),
            '{individual_copyright_registered_verified_marks}': str(data['B']['6']['copyrightIndividual']['ver_registeredMarks']),
            
            '{individual_copyright_granted_marks}': str(data['B']['6']['copyrightIndividual']['grantedCount'] * 50),
            '{individual_copyright_granted_verified_marks}': str(data['B']['6']['copyrightIndividual']['ver_grantedMarks']),
            '{individual_copyright_marks}': str(data['B']['6']['copyrightIndividual']['verified_marks']),
            
            # 7. Copyright Institute
            
            '{institute_copyright_registered_marks}': str(data['B']['7']['copyrightInstitute']['registeredCount'] * 40),
            '{institute_copyright_registered_verified_marks}': str(data['B']['7']['copyrightInstitute']['ver_registeredMarks']),
            
            '{institute_copyright_granted_marks}': str(data['B']['7']['copyrightInstitute']['grantedCount'] * 100),
            '{institute_copyright_granted_verified_marks}': str(data['B']['7']['copyrightInstitute']['ver_grantedMarks']),
            '{institute_copyright_marks}': str(data['B']['7']['copyrightInstitute']['verified_marks']),
            
            # 8-9. Patents (Individual and Institute)
            
            '{individual_patent_registered_marks}': str(data['B']['8']['patentIndividual']['registeredCount'] * 20),
            '{individual_patent_registered_verified_marks}': str(data['B']['8']['patentIndividual']['ver_registeredMarks']),
            
            '{individual_patent_published_marks}': str(data['B']['8']['patentIndividual']['publishedCount'] * 30),
            '{individual_patent_published_verified_marks}': str(data['B']['8']['patentIndividual']['ver_publishedMarks']),
            
            '{individual_granted_marks}': str(data['B']['8']['patentIndividual']['grantedCount'] * 50),
            '{individual_granted_verified_marks}': str(data['B']['8']['patentIndividual']['ver_grantedMarks']),
            
            '{individual_commercialized_marks}': str(data['B']['8']['patentIndividual']['commercializedCount'] * 100),
            '{individual_commercialized_verified_marks}': str(data['B']['8']['patentIndividual']['ver_commercializedMarks']),
            '{individual_patent_marks}': str(data['B']['8']['patentIndividual']['verified_marks']),
            
            #9
            
            '{college_patent_registered_marks}': str(data['B']['9']['patentInstitute']['registeredCount'] * 40),
            '{college_patent_registered_verified_marks}': str(data['B']['9']['patentInstitute']['ver_registeredMarks']),
            
            '{college_patent_published_marks}': str(data['B']['9']['patentInstitute']['publishedCount'] * 60),
            '{college_patent_published_verified_marks}': str(data['B']['9']['patentInstitute']['ver_publishedMarks']),
            
            '{college_granted_marks}': str(data['B']['9']['patentInstitute']['grantedCount'] * 100),
            '{college_granted_verified_marks}': str(data['B']['9']['patentInstitute']['ver_grantedMarks']),
            
            '{college_commercialized_marks}': str(data['B']['9']['patentInstitute']['commercializedCount'] * 200),
            '{college_commercialized_verified_marks}': str(data['B']['9']['patentInstitute']['ver_commercializedMarks']),
            '{college_patent_marks}': str(data['B']['9']['patentInstitute']['verified_marks']),
            '{patents_marks}': str(data['B']['8']['patentIndividual']['verified_marks'] + data['B']['9']['patentInstitute']['verified_marks']),
            
            # 10. Research Grants
            '{research_grants_amount}': str(data['B']['10']['researchGrants']['amount']),
            '{research_grants_marks}': str(math.floor(data['B']['10']['researchGrants']['amount'] / 200000) * 10),
            '{research_grants_verified_marks}': str(data['B']['10']['researchGrants']['ver_amountMarks']),
            
            # 11. Training Revenue
            '{training_amount}': str(data['B']['11']['trainingPrograms']['amount']),
            '{training_marks}': str(math.floor(data['B']['11']['trainingPrograms']['amount'] / 10000) * 5),
            '{training_verified_marks}': str(data['B']['11']['trainingPrograms']['ver_amountMarks']),
            
            # 12. Non-Research Grants
            '{nonresearch_grants_amount}': str(data['B']['12']['nonResearchGrants']['amount']),
            '{nonresearch_grants_marks}': str(math.floor(data['B']['12']['nonResearchGrants']['amount'] / 10000) * 5),
            '{nonresearch_grants_verified_marks}': str(data['B']['12']['nonResearchGrants']['ver_amountMarks']),
            
            # 13. Products
            
            '{commercialized_products_marks}': str(data['B']['13']['productDevelopment']['commercializedCount'] * 100),
            '{commercialized_products_verified_marks}': str(data['B']['13']['productDevelopment']['ver_commercializedMarks']),
            
            '{developed_products_marks}': str(data['B']['13']['productDevelopment']['developedCount'] * 40),
            '{developed_products_verified_marks}': str(data['B']['13']['productDevelopment']['ver_developedMarks']),
            
            '{poc_products_marks}': str(data['B']['13']['productDevelopment']['pocCount'] * 10),
            '{poc_products_verified_marks}': str(data['B']['13']['productDevelopment']['ver_pocMarks']),
            '{products_marks}': str(data['B']['13']['productDevelopment']['verified_marks']),
            
            # 14. Startup PCCOE
            '{startup_revenue_pccoe_amount}': str(data['B']['14']['startup']['revenueFiftyKCount']),
            '{startup_revenue_pccoe_marks}': str(data['B']['14']['startup']['revenueFiftyKCount'] * 100),
            '{startup_revenue_pccoe_verified_marks}': str(data['B']['14']['startup']['ver_revenueFiftyKMarks']),
            '{startup_funding_pccoe_amount}': str(data['B']['14']['startup']['fundsFiveLakhsCount']),
            '{startup_funding_pccoe_marks}': str(data['B']['14']['startup']['fundsFiveLakhsCount'] * 100),
            '{startup_funding_pccoe_verified_marks}': str(data['B']['14']['startup']['ver_fundsFiveLakhsMarks']),
            
            '{startup_products_marks}': str(data['B']['14']['startup']['productsCount'] * 40),
            '{startup_products_verified_marks}': str(data['B']['14']['startup']['ver_productsMarks']),
            
            '{startup_poc_marks}': str(data['B']['14']['startup']['pocCount'] * 10),
            '{startup_poc_verified_marks}': str(data['B']['14']['startup']['ver_pocMarks']),
            
            '{startup_registered_marks}': str(data['B']['14']['startup']['registeredCount'] * 5),
            '{startup_registered_verified_marks}': str(data['B']['14']['startup']['ver_registeredMarks']),
            '{startup_pccoe_marks}': str(data['B']['14']['startup']['verified_marks']),
            
            # 15. Awards
            
            '{international_awards_marks}': str(data['B']['15']['awardsAndFellowships']['internationalAwardsCount'] * 30),
            '{international_awards_verified_marks}': str(data['B']['15']['awardsAndFellowships']['ver_internationalAwardsMarks']),
            
            '{government_awards_marks}': str(data['B']['15']['awardsAndFellowships']['governmentAwardsCount'] * 20),
            '{government_awards_verified_marks}': str(data['B']['15']['awardsAndFellowships']['ver_governmentAwardsMarks']),
            
            '{national_awards_marks}': str(data['B']['15']['awardsAndFellowships']['nationalAwardsCount'] * 5),
            '{national_awards_verified_marks}': str(data['B']['15']['awardsAndFellowships']['ver_nationalAwardsMarks']),
            
            '{international_fellowship_marks}': str(data['B']['15']['awardsAndFellowships']['internationalFellowshipsCount'] * 50),
            '{international_fellowship_verified_marks}': str(data['B']['15']['awardsAndFellowships']['ver_internationalFellowshipsMarks']),
            
            '{national_fellowship_marks}': str(data['B']['15']['awardsAndFellowships']['nationalFellowshipsCount'] * 30),
            '{national_fellowship_verified_marks}': str(data['B']['15']['awardsAndFellowships']['ver_nationalFellowshipsMarks']),
            '{awards_marks}': str(data['B']['15']['awardsAndFellowships']['verified_marks']),
            
            # 16. Industry Interaction
            
            '{active_mou_marks}': str(data['B']['16']['industryInteraction']['moUsCount'] * 10),
            '{active_mou_verified_marks}': str(data['B']['16']['industryInteraction']['ver_moUsMarks']),
            
            '{lab_development_marks}': str(data['B']['16']['industryInteraction']['collaborationCount'] * 20),
            '{lab_development_verified_marks}': str(data['B']['16']['industryInteraction']['ver_collaborationMarks']),
            '{industry_interaction_marks}': str(data['B']['16']['industryInteraction']['verified_marks']),
            
            # 17. Industry Association
            
            '{internships_placements_marks}': str(data['B']['17']['internshipPlacement']['offersCount'] * 10),
            '{internships_placements_verified_marks}': str(data['B']['17']['internshipPlacement']['ver_offersMarks']),
            '{industry_association_marks}': str(data['B']['17']['internshipPlacement']['verified_marks']),
            
            # Total Section B
            '{B_total_marks}': str(b_total),
            '{section_b_total}': str(b_total_verified),
            '{Prof_B}': str(Prof_B),
            '{Assoc_B}': str(Assoc_B),
            '{Assis_B}': str(Assis_B),
            '{Prof_B_total_marks}': str(Prof_B_total_marks),
            '{Assoc_B_total_marks}': str(Assoc_B_total_marks),
            '{Assis_B_total_marks}': str(Assis_B_total_marks),
            '{Prof_B_total_verified}': str(Prof_B_total_verified),
            '{Assoc_B_total_verified}': str(Assoc_B_total_verified),
            '{Assis_B_total_verified}': str(Assis_B_total_verified),
            '{verf_committee_name}': verifier_name,
            
            # Section C placeholders
            '{Prof_qualification_marks}': str(Prof_qualification_marks),
            '{qualification_marks}': str(qualification_marks),
            '{training_attended_marks}': str(data['C']['2']['trainingAttended']['marks']),
            '{training_organized_marks}': str(data['C']['3']['trainingOrganized']['marks']),
            '{phd_guided_marks}': str(data['C']['4']['phdGuided']['marks']),
            '{section_c_total}': str(c_total),
            '{Prof_C}': str(Prof_C),
            '{Assoc_C}': str(Assoc_C),
            '{Assis_C}': str(Assis_C),
            '{Prof_C_total_marks}': str(Prof_C_total_marks),
            '{Assoc_C_total_marks}': str(Assoc_C_total_marks),
            '{Assis_C_total_marks}': str(Assis_C_total_marks),
            
            # New section D (Portfolio details)
            '{Department_portfolio}': "Not Applicable" if not data['D'].get('departmentLevelPortfolio') else data['D']['departmentLevelPortfolio'],
            '{Institute_Portfolio}': "Not Applicable" if not data['D'].get('instituteLevelPortfolio') else data['D']['instituteLevelPortfolio'],
            '{self_awarded_marks}': str(self_awarded_marks),
            '{hodMarks}': str(hod_marks),
            '{section_d_total}': str(total_marks_D),
            '{assDeanDeanMarks}' : str(assDeanDeanMarks),
            '{assDeanHODMarks}' : str(assDeanHODMarks),
            '{assTotalMarks}' : str(assTotalMarks),
            '{assSelfawardedmarks}' : str(assSelfawardedmarks),
            '{sumMarks_hod_dean}' : str(sumMarks_hod_dean),
            
            
            # Section E placeholders
            # '{section_E_total}': str(data['E']['total_marks']),
            '{section_E_total}': str(data.get('E', {}).get('total_marks', 0)),
            
            # Grand total
            '{total_for_C}' : str(round(data['C']['total_marks'])),
            '{total_for_B}' : str(round(data['B']['total_marks'])),
            '{total_for_A}' : str(round(data['A']['total_marks'])),
            '{total_for_D}' : str(round(data['D']['total_marks'])),
            '{total_for_B_verified}' : str(round(data['B']['final_verified_marks'])),
            '{grand_total}': str(round(data['grand_total']['grand_total'])),
            '{total_for_A_verified}' : str(round(data['A_verified_marks'])),
            '{total_for_C_verified}' : str(round(data['C_verified_marks'])),
            '{total_for_D_verified}' : str(round(data['D_verified_marks'])),
            '{total_for_E_verified}' : str(round(data['E_verified_marks'])),
            '{grand_verified_marks}': str(round(data['grand_verified_marks'])),
        })
        
        print('-----------after placeholders update-----------')
        # Replace placeholders in paragraphs and tables
        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for placeholder, value in placeholders.items():
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, value)
        
        return doc

    except Exception as e:
        raise Exception(f"Error in fill_template_document: {str(e)}")

# @app.route('/<department>/<user_id>/generate-doc', methods=['GET'])
# def generate_filled_document(department, user_id):
#     output_path = None
#     try:
#         # Get the department collection
#         collection = department_collections.get(department)
#         if collection is None:
#             return jsonify({"error": "Invalid department"}), 400

#         # Get user data directly from MongoDB
#         user_doc = collection.find_one({"_id": user_id})
#         if not user_doc:
#             return jsonify({"error": "User data not found"}), 404

#         # Create data structure
#         data = {
#             'A': user_doc.get('A', {}),
#             'B': user_doc.get('B', {}),
#             'C': user_doc.get('C', {})
#         }

#         # Generate document
#         doc = fill_template_document(data, user_id, department)
        
#         # Create a safe filename
#         safe_filename = secure_filename(f"filled_appraisal_{user_id}.docx")
        
#         # Save to a temporary directory
#         temp_dir = os.path.join(os.getcwd(), 'temp')
#         os.makedirs(temp_dir, exist_ok=True)
#         output_path = os.path.join(temp_dir, safe_filename)
        
#         # Save the document
#         doc.save(output_path)
        
#         # Send file
#         return send_file(
#             output_path,
#             as_attachment=True,
#             download_name=safe_filename,
#             mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
        
#     except Exception as e:
#         if output_path and os.path.exists(output_path):
#             try:
#                 os.remove(output_path)
#             except:
#                 pass
#         return jsonify({"error": str(e)}), 500
    
from docx2pdf import convert
import tempfile


@app.route('/<department>/<user_id>/generate-doc', methods=['GET'])
def generate_document(department, user_id):
    print(user_id)
    temp_docx = None
    output_path = None
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User data not found"}), 404

        # Check if document needs regeneration
        is_updated = user_doc.get('isUpdated', True)  # Default to True for backward compatibility
        existing_pdf = user_doc.get('appraisal_pdf')

        if not is_updated and existing_pdf:
            # Return existing PDF from GridFS
            try:
                file_id = ObjectId(existing_pdf['file_id'])
                grid_out = fs.get(file_id)
                
                return send_file(
                    io.BytesIO(grid_out.read()),
                    as_attachment=True,
                    download_name=existing_pdf['filename'],
                    mimetype='application/pdf'
                )
            except Exception as e:
                # If there's any error retrieving existing PDF, generate new one
                pass

        # Initialize COM for PDF generation
        pythoncom.CoInitialize()

        # Prepare data for document generation with proper grand_total structure
        user_doc = collection.find_one({"_id": user_id})
        grand_total_data = user_doc.get('grand_total', {'grand_total': 0, 'status': 'pending'})
        
        # Ensure grand_total is in correct format
        if isinstance(grand_total_data, (int, float)):
            grand_total_data = {'grand_total': float(grand_total_data), 'status': 'pending'}
        
        A_verified_marks = 0
        B_verified_marks = 0
        C_verified_marks = 0
        D_verified_marks = 0
        E_verified_marks = 0
        grand_verified_marks = 0
        
        if user_doc.get('grand_marks_A', {}).get('verified_marks'):
            A_verified_marks = user_doc['grand_marks_A']['verified_marks']
        if user_doc.get('grand_marks_B', {}).get('verified_marks'):
            B_verified_marks = user_doc['grand_marks_B']['verified_marks']
        if user_doc.get('grand_marks_C', {}).get('verified_marks'):
            C_verified_marks = user_doc['grand_marks_C']['verified_marks']
        if user_doc.get('grand_marks_D', {}).get('verified_marks'):
            D_verified_marks = user_doc['grand_marks_D']['verified_marks']
        if user_doc.get('grand_marks_E', {}).get('verified_marks'):
            E_verified_marks = user_doc['grand_marks_E']['verified_marks']
        if user_doc.get('grand_verified_marks'):
            grand_verified_marks = user_doc['grand_verified_marks']
        
        
        
        data = {
            'A': user_doc.get('A', {}),
            'B': user_doc.get('B', {}),
            'C': user_doc.get('C', {}),
            'D': user_doc.get('D', {}),
            'E': user_doc.get('E', {}),
            'grand_total': grand_total_data,
            'A_verified_marks' : A_verified_marks,
            'B_verified_marks' : B_verified_marks,
            'C_verified_marks' : C_verified_marks,
            'D_verified_marks' : D_verified_marks,
            'E_verified_marks' : E_verified_marks,
            'grand_verified_marks' : grand_verified_marks
            
        }

        doc = fill_template_document(data, user_id, department)
        
        # Create temporary directory
        temp_dir = os.path.join(os.getcwd(), 'temp')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Generate PDF
        safe_filename_docx = secure_filename(f"temp_{user_id}.docx")
        safe_filename = secure_filename(f"filled_appraisal_{user_id}.pdf")
        temp_docx = os.path.join(temp_dir, safe_filename_docx)
        output_path = os.path.join(temp_dir, safe_filename)
        
        # Save and convert to PDF
        doc.save(temp_docx)
        convert(temp_docx, output_path)
        
        # Store PDF in GridFS
        with open(output_path, 'rb') as pdf_file:
            file_id = fs.put(
                pdf_file,
                filename=safe_filename,
                user_id=user_id,
                department=department,
                content_type='application/pdf'
            )
        
        # Update user document with file reference and reset isUpdated flag
        collection.update_one(
            {"_id": user_id},
            {
                "$set": {
                    "appraisal_pdf": {
                        "file_id": str(file_id),
                        "filename": safe_filename,
                        "upload_date": datetime.now()
                    },
                    "isUpdated": False  # Reset flag after generating new PDF
                }
            }
        )
        
        # Send file
        return send_file(
            output_path,
            as_attachment=True,
            download_name=safe_filename,
            mimetype='application/pdf'
        )

    except Exception as e:
        # Cleanup on error
        for path in [temp_docx, output_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
        print(str(e))
        return jsonify({"error": str(e)}), 500
    finally:
        # Cleanup temporary files and uninitialize COM
        for path in [temp_docx, output_path]:
            if path and os.path.exists(path):
                try:
                    os.remove(path)
                except OSError:
                    pass
        pythoncom.CoUninitialize()

# Add a download endpoint to handle direct file downloads
@app.route('/download/<filename>')
def download_file(filename):
    try:
        return send_from_directory(
            os.path.join(os.getcwd(), 'temp'),
            filename,
            as_attachment=True
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 404

# Add this function to clean up old temporary files
def cleanup_temp_files():
    temp_dir = os.path.join(os.getcwd(), 'temp')
    if os.path.exists(temp_dir):
        for file in os.listdir(temp_dir):
            try:
                file_path = os.path.join(temp_dir, file)
                # Remove files older than 1 hour
                if os.path.getmtime(file_path) < time.time() - 3600:
                    os.remove(file_path)
            except:
                continue

# Add these imports at the top of your file
import time
from apscheduler.schedulers.background import BackgroundScheduler
import pythoncom
from datetime import datetime

# Add this after your app initialization
scheduler = BackgroundScheduler()
scheduler.add_job(func=cleanup_temp_files, trigger="interval", minutes=60)
scheduler.start()

@app.route('/<department>/<user_id>/download/<format>', methods=['GET'])
def get_stored_document(department, user_id, format):
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400
            
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404
            
        # Get file reference based on format
        file_ref = user_doc.get(f"appraisal_{format}")
        if not file_ref:
            return jsonify({"error": "Document not found"}), 404
            
        # Get file from GridFS
        file_id = ObjectId(file_ref['file_id'])
        grid_out = fs.get(file_id)
        
        # Return the file
        return send_file(
            io.BytesIO(grid_out.read()),
            as_attachment=True,
            download_name=file_ref['filename'],
            mimetype=grid_out.content_type
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Add after your app initialization and before running the app
verification_bp = create_verification_blueprint(mongo_fdw, db_users, department_collections)
app.register_blueprint(verification_bp)
# Add this line after creating the Flask app
app.register_blueprint(faculty_list)

# Add this line with your other blueprint registrations
app.register_blueprint(forgot_password)

@app.route('/<department>/<user_id>/get-status', methods=['GET'])
def get_status(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400
        user_doc  = collection.find_one({"_id" : user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404
        current_status = user_doc.get("status","pending")
        return jsonify({
                "message": "Form submitted successfully",
                "status": f"{current_status}"
            }), 200
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Add these status change endpoints after your existing routes
@app.route('/<department>/<user_id>/submit-form', methods=['POST'])
def submit_form(department, user_id):
    """Changes status from 'pending' to 'verification_pending' when form is submitted"""
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Get current document and check status
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        current_status = user_doc.get('status', 'pending')
        if current_status != 'pending':
            return jsonify({
                "error": "Invalid status transition",
                "message": "Form must be in pending status to submit"
            }), 400

        # Update status
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {"status": "Portfolio_Mark_pending"}}
        )

        if result.modified_count > 0:
            return jsonify({
                "message": "Form submitted successfully",
                "new_status": "Portfolio_Mark_pending"
            }), 200
        return jsonify({"error": "No changes made"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500



@app.route('/<department>/<user_id>/hod-mark-given', methods=['POST'])
def hod_mark_given(department, user_id):
    """Changes status from 'Portfolio_Mark_pending' to 'Portfolio_Mark_Dean_pending' after HOD assigns marks"""
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Get current document and check status
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        current_status = user_doc.get('status', 'pending')
        if current_status != 'Portfolio_Mark_pending':
            return jsonify({
                "error": "Invalid status transition",
                "message": "Form must be in Portfolio_Mark_pending status to proceed"
            }), 400

        # Update status to indicate Dean marks are pending
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {"status": "Portfolio_Mark_Dean_pending"}}
        )

        if result.modified_count > 0:
            return jsonify({
                "message": "HOD portfolio marks assigned successfully, awaiting Dean review",
                "new_status": "Portfolio_Mark_Dean_pending"
            }), 200
        return jsonify({"error": "No changes made"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/portfolio-given', methods=['POST'])
def portfolio_given(department, user_id):
    """Changes status from 'Portfolio_Mark_pending' or 'Portfolio_Mark_Dean_pending' to 'verification_pending' when portfolio marks are assigned"""
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Get current document and check status
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        current_status = user_doc.get('status', 'pending')
        if current_status != 'Portfolio_Mark_pending' and current_status != 'Portfolio_Mark_Dean_pending':
            return jsonify({
                "error": "Invalid status transition",
                "message": "Form must be in Portfolio_Mark_pending or Portfolio_Mark_Dean_pending status to proceed"
            }), 400

        # Update status
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {"status": "verification_pending"}}
        )

        if result.modified_count > 0:
            return jsonify({
                "message": "Portfolio marks assigned successfully",
                "new_status": "verification_pending"
            }), 200
        return jsonify({"error": "No changes made"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<verifier_id>/<user_id>/verify-research', methods=['POST'])
def verify_research(department,verifier_id, user_id):
    """Changes status from 'verification_pending' to 'authority_verification_pending'"""
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Get current document and check status
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        current_status = user_doc.get('status')
        if current_status != 'verification_pending':
            return jsonify({
                "error": "Invalid status transition",
                "message": "Form must be in verification_pending status"
            }), 400

        # Update status
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {"status": "authority_verification_pending"}}
        )
        
        committee_head = db_users.find_one({"_id": verifier_id})
        if not committee_head:
            return jsonify({"error": "Committee head not found"}), 404

        # Check if the user is in verification panel
        if not committee_head.get("isInVerificationPanel", False):
            return jsonify({"error": "User is not authorized to approve"}), 403

        # Update the isApproved status in facultyToVerify array
        result_isVerified = db_users.update_one(
            {
                "_id": verifier_id,
                f"facultyToVerify.{department}._id": user_id
            },
            {
                "$set": {
                    f"facultyToVerify.{department}.$.isApproved": True
                }
            }
        )
        if result_isVerified.modified_count <= 0:
            return jsonify({"error": "Faculty not found or already approved"}), 400

        if result.modified_count > 0:
            return jsonify({
                "message": "Research verification completed",
                "new_status": "authority_verification_pending",
                "department": department,
                "faculty_id": user_id,
                "verifier_id": verifier_id
            }), 200
        return jsonify({"error": "No changes made"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/verify-authority', methods=['POST'])
def verify_authority(department, user_id):
    """Changes status from 'authority_verification_pending' to 'Interaction_pending'"""
    try:
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Get current document and check status
        user_doc = collection.find_one({"_id": user_id})
        if not user_doc:
            return jsonify({"error": "User not found"}), 404

        current_status = user_doc.get('status')
        if current_status != 'verified':
            return jsonify({
                "error": "Invalid status transition",
                "message": "Form must be in verified status"
            }), 400

        # Update status
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {"status": "Interaction_pending"}}
        )
        if result.modified_count > 0:
            return jsonify({
                "message": "Authority verification completed",
                "new_status": "Interaction_pending"
            }), 200
        return jsonify({"error": "No changes made"}), 400

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    
@app.route('/<department>/send-to-director', methods=['POST'])
def send_to_director(department):
    """Changes status from 'Done' to 'SentToDirector' for multiple users at once"""
    try:
        data = request.get_json()
        if not data or 'user_ids' not in data or not isinstance(data['user_ids'], list):
            return jsonify({"error": "Missing required field: user_ids (array)"}), 400
        
        user_ids = data['user_ids']
        if not user_ids:
            return jsonify({"error": "Empty user_ids array provided"}), 400

        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400

        # Find users with 'Done' status
        valid_users = collection.find({
            "_id": {"$in": user_ids},
            "status": "done"
        })
        
        valid_user_ids = [user["_id"] for user in valid_users]
        
        if not valid_user_ids:
            return jsonify({
                "error": "No valid users found",
                "message": "No users with 'Done' status found among the provided IDs"
            }), 404

        # Update status for all valid users
        result = collection.update_many(
            {
                "_id": {"$in": valid_user_ids},
                "status": "done"
            },
            {"$set": {"status": "SentToDirector"}}
        )

        # Check results
        success_count = result.modified_count
        skipped_ids = [user_id for user_id in user_ids if user_id not in valid_user_ids]

        return jsonify({
            "message": f"Successfully sent {success_count} form(s) to director",
            "successful_ids": valid_user_ids,
            "skipped_ids": skipped_ids,
            "new_status": "SentToDirector"
        }), 200

    except Exception as e:
        return jsonify({"error": str(e)}), 500
    


#To get all the deans
# Add this after your existing routes

@app.route('/deans', methods=['GET'])
def get_deans():
    """Get all users with role 'Dean'"""
    try:
        # Find users with role 'Dean'
        deans = db_users.find({"role": "Dean"})
        
        if not deans:
            return jsonify({
                "message": "No deans found",
                "data": []
            }), 404
            
        # Convert cursor to list and return
        dean_list = list(deans)
        
        return jsonify({
            "message": "Deans retrieved successfully",
            "data": json.loads(dumps(dean_list))
        }), 200
            
    except Exception as e:
        return jsonify({
            "error": "Failed to retrieve deans",
            "message": str(e)
        }), 500



from dean_associates import dean_associates

# Add this line with your other blueprint registrations
app.register_blueprint(dean_associates)

from externals import externals

# Add this line with your other blueprint registrations
app.register_blueprint(externals)

@app.route('/<department>/<user_id>/E', methods=['POST'])
def handle_post_E(department, user_id):
    try:
        data = request.get_json()
        if not data:
            return jsonify({"error": "Invalid JSON data"}), 400

        # Get department collection
        collection = department_collections.get(department)
        if collection is None:
            return jsonify({"error": "Invalid department"}), 400
        
        # Verify user exists in department
        lookup = collection.find_one({"_id": "lookup"}).get("data")
        if lookup is None:
            return jsonify({"error": "Invalid department"}), 400
        user = lookup.get(user_id)
        if user is None:
            return jsonify({"error": "Invalid user"}), 400

        # Initialize default structure for section E
        if 'E' not in data:
            data['E'] = {}

        # Validate required fields
        if 'total_marks' not in data['E']:
            return jsonify({"error": "Missing required field: total_marks"}), 400
        if 'bullet_points' not in data['E']:
            return jsonify({"error": "Missing required field: bullet_points"}), 400

        # Create section E structure with default values
        section_E = {
            'total_marks': data['E']['total_marks'],
            'bullet_points': data['E']['bullet_points'],
            'verified_marks': 0,  # Default verified marks
            'isVerified': False  # Verification status
        }

        # Update the document
        result = collection.update_one(
            {"_id": user_id},
            {"$set": {
                "E": section_E,
                "isUpdated": True,
                "status": "pending"
            }},
            upsert=True
        )

        # Get updated document and calculate grand total
        updated_doc = collection.find_one({"_id": user_id})
        calculated_data = calculate_grand_total(updated_doc)

        # Update grand total and status
        collection.update_one(
            {"_id": user_id},
            {"$set": {
                "grand_total": calculated_data['grand_total'],
                "status": calculated_data['status']
            }}
        )

        return jsonify({
            "message": "Data updated successfully" if result.matched_count > 0 else "Data inserted successfully",
            "grand_total": calculated_data['grand_total'],
            "status": calculated_data['status']
        }), 200

    except Exception as e:
        print(f"Error updating section E: {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/<department>/<user_id>/E', methods=['GET'])
def get_section_E(department, user_id):
    try:
        collection = department_collections.get(department)
        if collection is not None:
            user = collection.find_one({"_id": user_id})
            if user:
                return jsonify(user.get("E", {
                    'total_marks': 0,
                    'bullet_points': [],
                    'verified_marks': 0,
                    'isVerified': False
                }))
            return jsonify({"error": "User not found"}), 404
        return jsonify({"error": "Invalid department"}), 400
    except Exception as e:
        print(f"Error retrieving section E: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
